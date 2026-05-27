"""
Convert תיק_פרויקט.md  →  תיק_פרויקט.docx
Plain style: black text, no decorative colors.
RTL Hebrew, Heading 1/2/3/4, tables, code blocks, bullets.
"""
import re, pathlib
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(__file__).parent
SRC  = ROOT / "תיק_פרויקט.md"
OUT  = ROOT / "תיק_פרויקט.docx"

FONT_HEB  = "Arial"
FONT_MONO = "Courier New"
C_BLACK   = RGBColor(0x00, 0x00, 0x00)
C_DARK    = RGBColor(0x1a, 0x1a, 0x1a)
C_GRAY    = RGBColor(0x55, 0x55, 0x55)
C_CODE_BG = "F2F2F2"
C_TH_BG   = "D0D0D0"
C_ALT_BG  = "F5F5F5"


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_rtl_para(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.insert(0, bidi)
    for old in pPr.findall(qn("w:jc")):
        pPr.remove(old)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)


def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)
    rPr.append(OxmlElement("w:cs"))


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def set_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"),  str(after))
    if line:
        sp.set(qn("w:line"),     str(line))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def set_indent(para, right=None, left=None):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    if right is not None: ind.set(qn("w:right"), str(right))
    if left  is not None: ind.set(qn("w:left"),  str(left))
    pPr.append(ind)


def set_para_border_bottom(para, size="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pPr.append(pBdr)


# ── Styles ─────────────────────────────────────────────────────────────────────

def configure_styles(doc):
    s = doc.styles

    n = s["Normal"]
    n.font.name = FONT_HEB
    n.font.size = Pt(11)
    n.font.color.rgb = C_DARK
    n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for level, size, bold in [(1, 18, True), (2, 14, True), (3, 12, True), (4, 11, True)]:
        h = s[f"Heading {level}"]
        h.font.name = FONT_HEB
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = C_BLACK
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for style_name in ("List Bullet", "List Number"):
        ls = s[style_name]
        ls.font.name = FONT_HEB
        ls.font.size = Pt(11)
        ls.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ── Inline markup ──────────────────────────────────────────────────────────────

def parse_inline(para, text, base_size=11, mono=False):
    pattern = re.compile(
        r"\*\*(.+?)\*\*"
        r"|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)"
        r"|`(.+?)`"
        r"|\[(.+?)\]\((.+?)\)"
        r"|(‎)"
    )
    pos = 0
    for m in pattern.finditer(text):
        chunk = text[pos:m.start()]
        if chunk:
            r = para.add_run(chunk)
            r.font.name = FONT_MONO if mono else FONT_HEB
            r.font.size = Pt(base_size)
            set_rtl_run(r)

        if m.group(1):
            r = para.add_run(m.group(1))
            r.bold = True
            r.font.name = FONT_HEB
            r.font.size = Pt(base_size)
            set_rtl_run(r)
        elif m.group(2):
            r = para.add_run(m.group(2))
            r.italic = True
            r.font.name = FONT_HEB
            r.font.size = Pt(base_size)
            set_rtl_run(r)
        elif m.group(3):
            r = para.add_run(m.group(3))
            r.font.name = FONT_MONO
            r.font.size = Pt(9)
            set_rtl_run(r)
        elif m.group(4):
            r = para.add_run(m.group(4))
            r.font.name = FONT_HEB
            r.font.size = Pt(base_size)
            r.underline = True
            set_rtl_run(r)
        pos = m.end()

    tail = text[pos:]
    if tail:
        r = para.add_run(tail)
        r.font.name = FONT_MONO if mono else FONT_HEB
        r.font.size = Pt(base_size)
        set_rtl_run(r)


# ── Builders ───────────────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", re.sub(r"`(.+?)`", r"\1", text))
    para = doc.add_paragraph(style=f"Heading {level}")
    set_rtl_para(para)
    if level == 1:
        set_spacing(para, before=240, after=80)
        set_para_border_bottom(para, "8")
    elif level == 2:
        set_spacing(para, before=160, after=60)
        set_para_border_bottom(para, "4")
    elif level == 3:
        set_spacing(para, before=100, after=40)
    else:
        set_spacing(para, before=60, after=20)
    r = para.add_run(clean)
    r.font.name = FONT_HEB
    set_rtl_run(r)


def add_body(doc, text):
    if not text.strip():
        return
    para = doc.add_paragraph(style="Normal")
    set_rtl_para(para)
    set_spacing(para, before=20, after=60, line=276)
    parse_inline(para, text)


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    set_rtl_para(para)
    if level > 0:
        set_indent(para, right=level * 360)
    set_spacing(para, before=20, after=20, line=252)
    parse_inline(para, text)


def add_numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    set_rtl_para(para)
    set_spacing(para, before=20, after=20, line=252)
    parse_inline(para, text)


def add_code_block(doc, lines):
    if not lines:
        return
    tbl = doc.add_table(rows=len(lines), cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, line in enumerate(lines):
        cell = tbl.rows[i].cells[0]
        set_cell_bg(cell, C_CODE_BG)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(para, before=0, after=0, line=220)
        r = para.add_run(line if line else " ")
        r.font.name = FONT_MONO
        r.font.size = Pt(8.5)
        r.font.color.rgb = C_DARK
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn("w:bidi")):
            pPr.remove(old)
    for row in tbl.rows:
        row.cells[0].width = Cm(17)
    doc.add_paragraph()


def add_table(doc, rows):
    if not rows:
        return
    col_n = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_n)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_width = Cm(17.0 / col_n)

    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri]
        for ci in range(col_n):
            cell = row.cells[ci]
            cell.width = col_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                set_cell_bg(cell, C_TH_BG)
            elif ri % 2 == 1:
                set_cell_bg(cell, C_ALT_BG)
            para = cell.paragraphs[0]
            set_rtl_para(para)
            set_spacing(para, before=40, after=40)
            txt = row_data[ci] if ci < len(row_data) else ""
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1",
                   re.sub(r"`(.+?)`", r"\1", txt.replace("‎", "")))
            r = para.add_run(clean)
            r.font.name = FONT_HEB
            r.font.size = Pt(9.5)
            r.font.color.rgb = C_DARK
            if ri == 0:
                r.bold = True
            set_rtl_run(r)
    doc.add_paragraph()


def add_blockquote(doc, text):
    para = doc.add_paragraph(style="Normal")
    set_rtl_para(para)
    set_indent(para, right=360, left=360)
    set_spacing(para, before=60, after=60)
    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", re.sub(r"`(.+?)`", r"\1", text))
    r = para.add_run(clean)
    r.font.name = FONT_HEB
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = C_GRAY
    set_rtl_run(r)


def add_hr(doc):
    para = doc.add_paragraph()
    set_para_border_bottom(para, "4")
    set_spacing(para, before=60, after=60)


def setup_header_footer(doc):
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = section.right_margin = Cm(2.5)
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.different_first_page_header_footer = False

    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    set_rtl_para(hp)
    r = hp.add_run("בלאק‑ג'ק רב‑משתתפים  |  תיק פרויקט")
    r.font.name  = FONT_HEB
    r.font.size  = Pt(8)
    r.font.color.rgb = C_GRAY
    set_rtl_run(r)
    set_para_border_bottom(hp, "4")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_xml = (
        '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:instr=" PAGE \\* MERGEFORMAT ">'
        '<w:r><w:rPr><w:sz w:val="16"/></w:rPr><w:t>1</w:t></w:r>'
        '</w:fldSimple>'
    )
    import lxml.etree as etree
    fp._p.append(etree.fromstring(fld_xml))


# ── Parser ─────────────────────────────────────────────────────────────────────

def parse_md(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    in_code  = False
    code_buf = []
    in_table = False
    table_rows = []
    first_h1 = True

    def flush_table():
        nonlocal in_table, table_rows
        if table_rows:
            add_table(doc, table_rows)
        in_table = False
        table_rows = []

    def flush_code():
        nonlocal in_code, code_buf
        add_code_block(doc, code_buf)
        in_code  = False
        code_buf = []

    while i < len(lines):
        raw     = lines[i]
        stripped = raw.rstrip()

        if stripped.startswith("```"):
            if in_table: flush_table()
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                flush_code()
            i += 1
            continue

        if in_code:
            code_buf.append(raw.rstrip("\n"))
            i += 1
            continue

        if stripped.startswith("|"):
            in_table = True
            cells = [c.strip() for c in stripped.split("|")
                     if c.strip() and c.strip() != ""]
            if not all(re.match(r"^[-: ]+$", c) for c in cells):
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:], 4)
            i += 1; continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 3)
            i += 1; continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 2)
            i += 1; continue
        if stripped.startswith("# "):
            txt = stripped[2:]
            if first_h1:
                first_h1 = False
                add_heading(doc, txt, 1)
            else:
                add_heading(doc, txt, 1)
            i += 1; continue

        if re.match(r"^-{3,}$", stripped):
            add_hr(doc)
            i += 1; continue

        if stripped.startswith("> "):
            add_blockquote(doc, stripped[2:])
            i += 1; continue

        m = re.match(r"^(\s*)[*\-]\s+(.+)", stripped)
        if m:
            level = len(m.group(1)) // 2
            add_bullet(doc, m.group(2), level)
            i += 1; continue

        m = re.match(r"^\s*\d+\.\s+(.+)", stripped)
        if m:
            add_numbered(doc, m.group(1))
            i += 1; continue

        add_body(doc, stripped)
        i += 1

    if in_table: flush_table()
    if in_code and code_buf: flush_code()


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    md_text = SRC.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    setup_header_footer(doc)
    parse_md(doc, md_text)
    doc.save(str(OUT))
    print(f"Saved -> {OUT}  ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
